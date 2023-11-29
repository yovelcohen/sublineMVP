import re
from datetime import timedelta

from docx import Document
from srt import Subtitle


class SRTBlock(Subtitle):
    def __init__(self, index, start, end, content, speaker=None, region=None, style=None, num_tokens=None,
                 translation=None):
        super().__init__(index, start, end, content)
        self.speaker = speaker
        self.region = region
        self.style = style
        self.num_tokens = num_tokens
        self.translation = translation

    def __repr__(self):
        return f'{self.index}\n{self.content}'


def parse_time(time_str):
    """ Convert a time string to a timedelta object. """
    hours, minutes, seconds = map(int, time_str.split(':'))
    return timedelta(hours=hours, minutes=minutes, seconds=seconds)


def parse_subtitle_file(file_content):
    """ Parse the subtitle file content into SRTBlock objects. """
    subtitles = []
    blocks = file_content.split('\n\n')
    for block in blocks:
        lines = block.strip().split('\n')
        if len(lines) < 3:
            continue

        index_line = lines[0]
        time_line = lines[1]
        content_lines = lines[2:]

        # Extract start and end times
        start_str, _, end_str = time_line.partition('-->')
        start = parse_time(start_str.strip())
        end = parse_time(end_str.strip())

        # Combine content lines
        content = ' '.join(content_lines)

        # Create a new SRTBlock object
        subtitles.append(SRTBlock(index=index_line, start=start, end=end, content=content))

    return subtitles


def parse_subtitles_doc(doc: Document):
    rows = []
    assert len(doc.tables) == 1
    table = doc.tables[0]
    for row in table.rows:
        timestamp = row.cells[0].text
        text = row.cells[2].text if row.cells[1] == '' else row.cells[1].text
        index, start, end, content = row.cells
        rows.append(SRTBlock(index=index, start=start, end=end, content=content))


if __name__ == '__main__':
    p = '/Users/yovel.c/PycharmProjects/services/sublineStreamlit/DREAM KITCHENS - S3 Ep01 Script.docx'
    # with open(p, 'r') as f:
    #     srt = f.read()
    srt = Document(p)
    ret = parse_subtitle_file(srt)
    print(ret)
